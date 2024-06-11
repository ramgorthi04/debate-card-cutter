from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.oxml.ns import qn
import re

# Sample paragraph
paragraph_text = "This is the first sentence. This is the second sentence. This is the third sentence."

# Create a new Document
doc = Document()

# Add a paragraph to the document
p = doc.add_paragraph()

# Split the paragraph into sentences
sentences = re.split(r'(?<=[.!?])\s+', paragraph_text)

# Underline the first two sentences and highlight the phrase "the second sentence" in green
for i, sentence in enumerate(sentences):
    if "the second sentence" in sentence:
        parts = sentence.split("the second sentence")
        run1 = p.add_run(parts[0])
        run2 = p.add_run("the ")
        run3 = p.add_run("second")
        run4 = p.add_run(" sentence")
        run5 = p.add_run(parts[1] if len(parts) > 1 else "")
        
        # Highlight and bold 'second'
        run3.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        run3.bold = True
        
        # Underline the first two sentences
        if i < 2:
            run1.underline = True
            run2.underline = True
            run3.underline = True
            run4.underline = True
            run5.underline = True

        # Set font to Calibri and size to 11
        for run in [run1, run2, run3, run4, run5]:
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(11)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            r.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
            r.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
            r.rPr.rFonts.set(qn('w:cs'), 'Calibri')
    else:
        run = p.add_run(sentence + " ")
        if i < 2:
            run.underline = True

        # Set font to Calibri and size to 11
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(11)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        r.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        r.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        r.rPr.rFonts.set(qn('w:cs'), 'Calibri')

# Save the document
doc.save('underlined_highlighted_bold.docx')

print("Document created successfully.")

def add_run_to_paragraph(p, run, underline, bold, highlight):
    run = p.add_run(run)
    run.underline = underline
    run.bold = bold
    font = run.font
    font.name = 'Calibri'
    if underline:
        font.size = Pt(11)
    else:
        font.size = Pt(8)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        r.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        r.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        r.rPr.rFonts.set(qn('w:cs'), 'Calibri')
            
        if highlight:
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
