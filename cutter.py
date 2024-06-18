import json
import re
import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from scraper_test import *
from tools import *
import time
from datetime import date
import argparse
from research import *


GPT_MODEL = "gpt-4o-2024-05-13"
# Configure logging
logging.basicConfig(filename='output.log', level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')

def llm_get_text_to_underline(body_text: str, tag_line: str) -> list:
    """
    Given the following article text, returns a list of sentences/phrases to be underlined
    """
    logging.info('Calling llm_get_text_to_underline with body_text and tag_line.')
    
    sentence_endings = re.compile(r'[.!?]')
    # Find all matches of the sentence-ending punctuation
    sentences = sentence_endings.findall(body_text)
    
    # The number of sentences is the number of matches
    num_sentences = len(sentences)
    
    # Underline 50% of the sentences
    num_sentences_to_underline = int(num_sentences * 0.3)
    prompt = f"""Given the following text, return a list of the sentences and phrases that support this argument: {tag_line}. The sentences must be IN THE ORDER they appear in the text. Do not include ANY words that are not present in the article. 
    
    Return your response as a JSON in this form ["FIRST SENTENCE/PHRASE", "SECOND SENTENCE/PHRASE", "THIRD SENTENCE/PHRASE", ...]
    
    Article: {body_text}"""
    underlined_text_list = get_gpt_response(prompt, gpt_model=GPT_MODEL, json_mode=True)
    spliced_list = underlined_text_list[7:]
    spliced_list = spliced_list[:-3]
    
    # Try to load the spliced list to JSON
    try:
        spliced_list = json.loads(spliced_list)
        underlined_text_list = spliced_list
        logging.info('Underlined text list successfully parsed.')
    except json.JSONDecodeError as e:
        logging.error(f"Error parsing underlined text list: {e}")
        logging.error(underlined_text_list)
    
        # Now try to parse with regex
        list_pattern = re.search(r'\[.*?\]', str, re.DOTALL)
        if list_pattern:
            # Extract the matched string
            list_string = list_pattern.group()
            try:
                # Convert the string to a list using json.loads
                underlined_text_list = json.loads(list_string)
                logging.info('Underlined text list successfully parsed.')
            except json.JSONDecodeError as e:
                logging.error(f"Error parsing underlined text list: {e}")
                logging.error(underlined_text_list)
                return []
        else:
            print("No list found in the input string.")
            logging.error(f"Error parsing underlined text list")
            logging.error(underlined_text_list)
            return []

    logging.info(f"Underlined text list: {underlined_text_list}")
    return underlined_text_list

def llm_get_text_to_bold(underlined_text: str, num_words: int, tag_line: str) -> list:
    """
    Given underlined text, picks which words to bold. Returns a list of words that should be bolded.
    """
    logging.info('Calling llm_get_text_to_bold with underlined_text and tag_line.')

    # Bold ~10% of the words
    num_words_to_bold = int(num_words * 0.1)
    prompt = f"""Given the following article text, return a list of the individual words in the order they appear that should be emphasized in speech to support this argument: {tag_line}. Aim to have at least {num_words_to_bold} words bolded.
    
    Return your response as a JSON in this form ["word 1", "word 2", "word 3", ...]
    Article text: {underlined_text}"""
    
    words_list = get_gpt_response(prompt, gpt_model=GPT_MODEL, json_mode=True)
    words_list = words_list[7:]
    words_list = words_list[:-3]
    logging.info(f"Bolded words list before parsing: {words_list}")

    # Convert list to a set of words
    bolded_list = words_list.strip('[]').split(', ')
    # Check if bolded_list is not of type list, if so load it to list
    if not isinstance(bolded_list, list):
        bolded_list = json.loads(bolded_list)

    # Join the list elements into a single string
    joined_string = ' '.join(bolded_list)
    # Remove unwanted characters
    cleaned_string = re.sub(r'[\"\[\]\n]', '', joined_string)
    # Split the cleaned string into individual words
    bolded_list = cleaned_string.split()
    
    logging.info(f"Bolded words list: {bolded_list}")
    return bolded_list

def llm_get_text_to_highlight(underlined_text: str, tag_line: str) -> list:
    """
    Given underlined text, picks which words to highlight. Returns a list of phrases/sentences that should be highlighted.
    """
    logging.info('Calling llm_get_text_to_highlight with underlined_text and tag_line.')

    prompt = f"""Given the following article text, return a list of phrases/sentences in the order they appear to be read out loud as evidence to support this argument {tag_line}. Prioritize data and studies. Only about 40-50 words should be highlighted.
    Return response as a JSON of this form ['PHRASE/SENTENCE 1', 'PHRASE/SENTENCE 2', 'PHRASE/SENTENCE 3', ...]
    Article text: {underlined_text}"""
    
    highlighted_text = get_gpt_response(prompt, gpt_model=GPT_MODEL, json_mode=True)
    highlighted_text = highlighted_text[7:]
    highlighted_text = highlighted_text[:-3]

    logging.info(f"Highlighted text list: {highlighted_text}")
    logging.info("Processing highlighted text")
    
    # Regular expression to match words and symbols, excluding brackets, double quotes, and newlines, but preserving single quotes
    pattern = r"[^\[\]\"\n]+"

    # Find all matches
    matches = re.findall(pattern, highlighted_text)

    # Strip leading/trailing whitespace from each match and filter out empty strings
    matches = [match.strip() for match in matches if match.strip()]
    # Delete "," if it constitutes an entire list item
    cleaned_matches = []
    for match in matches:
        if match != ",":
            cleaned_matches.append(match)
    return cleaned_matches

def llm_cut_article(title: str, author: str, date: str, body_text: str, tag_line: str):
    """
    Given title, author, date, article text, and a tag line, returns a list of sentences to be underlined, bolded, and highlighted based on the argument in the tag line.
    Returns Tuple[list, list, list] or Tuple[None, None, None]
    """
    logging.info(f'Calling llm_cut_article with title: {title}, author: {author}, date: {date}, tag_line: {tag_line}.')

    # Get underlined sentences
    underlined_text_list = llm_get_text_to_underline(body_text, tag_line)
    if underlined_text_list == []:
        print("Failed to cut card. Please try again.")
        return None, None, None
    # Check case where LLM puts tag into underlined_text_list
    # Remove last character from tag if it is a period
    if tag_line[-1] == '.':
        tag_line = tag_line[:-1]
    if tag_line in underlined_text_list[0]:
        # Drop first element
        print("Removing tag from underlined list")
        underlined_text_list = underlined_text_list[1:]
    # Flatten list into single string and count number of words underlined
    underlined_text = ""
    num_words = 0
    for text in underlined_text_list:
        num_words += len(text.split())
        # Add text to underlined_text
        underlined_text += " " + text + ""
    
    # Out of underlined sentences, pick words to bold
    bolded_text_list = llm_get_text_to_bold(underlined_text, num_words, tag_line)
    # Out of underlined sentences, pick phrases/sentences to read out-loud
    highlighted_text_list = llm_get_text_to_highlight(underlined_text, tag_line)
    
    logging.info(f'Underlined text list: {underlined_text_list}')
    logging.info(f'Bolded text list: {bolded_text_list}')
    logging.info(f'Highlighted text list: {highlighted_text_list}')

    return underlined_text_list, bolded_text_list, highlighted_text_list

def split_string(text: str) -> list:
    """
    Returns list of words, spaces, punctuation, numbers, and paragraph breaks in string
    """
    pattern = r'\w+|[^\w\s]|\s+|\n\n+'
    
    # Split the text according to the pattern
    elements = re.findall(pattern, text)
    
    # # Turn each word to lower-case
    # elements = [element.lower() for element in elements]
    
    return elements

def add_run_to_paragraph(p, run, underline, bold, highlight):
    run = p.add_run(run)
    run.underline = underline
    run.bold = bold
    font = run.font
    font.name = 'Calibri'
    if underline:
        font.size = Pt(11)
    else:
        font.size = Pt(8)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        r.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        r.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        r.rPr.rFonts.set(qn('w:cs'), 'Calibri')
            
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

def splice_card_top_bottom(formatting_table: list) -> list:
    # Splice the formatting table by removing paragraphs at beginning/end with no underlining
    first_relevant_paragraph_index = 0
    for index, element in enumerate(formatting_table):
        # Check if the string contains '\n'
        if ('\n' in element[0]):
            first_relevant_paragraph_index = index
        if (element[1] == True):
            # An element has been underlined, so this paragraph must be included
            break
    last_relevant_paragraph_index = len(formatting_table) - 1
    for index, element in enumerate(formatting_table[::-1]):
        if ('\n' in element[0]):
            last_relevant_paragraph_index = len(formatting_table) - index - 1
        if (element[1] == True):
            # An element has been underlined, so this paragraph must be included
            break
    # Splice formatting table by removing paragraphs at beginning/end with no underlining
    formatting_table = formatting_table[first_relevant_paragraph_index:last_relevant_paragraph_index]
        
    return formatting_table

def get_nearest_words_left_right(list, index):
    """Given a list and index in the list, return the nearest word to the left and right of the index in the list."""

    # Find nearest word to the left
    left_iterator = index - 1
    while (left_iterator > 0):
        if list[left_iterator].isalpha():
            break
        left_iterator -= 1
    # Find nearest word to the right
    right_iterator = index + 1
    while (right_iterator < len(list)):
        if list[right_iterator].isalpha():
            break
        right_iterator += 1
    # Return the nearest words
    if left_iterator >= 0:
        left_word = list[left_iterator]
    else:
        left_word = None
    if right_iterator < len(list):
        right_word = list[right_iterator]
    else:
        right_word = None
    return left_word, right_word

def nearest_word_neighbors_equal(list_one, list_two, index_one, index_two) -> bool:
    """Given two lists and index in the list, check if the nearest word to the left and right of the indices in both lists are equal. If one is equal, return True. Otherwise, return False."""
    
    list_one_left, list_one_right = get_nearest_words_left_right(list_one, index_one)
    list_two_left, list_two_right = get_nearest_words_left_right(list_two, index_two)
    
    # Check that at least one set of words is equal:
    logging.info(f'Current word is {list_one[index_one]} which should equal {list_two[index_two]}')
    logging.info(f'List one left: {list_one_left}, List one right: {list_one_right}')
    logging.info(f'List two left: {list_two_left}, List two right: {list_two_right}')
    if list_one_left == list_two_left or list_one_right == list_two_right:
        logging.info(f'Nearest words are equal')
        return True
    else:
        logging.info(f'Nearest words are not equal')
        return False
    
def format_card(article_info: dict, tag: str, all_text: str, underlined_text_list: list, bolded_text_list: list, highlighted_text_list: list) -> None:
    """Produces word document. Underlines, bolds, and highlights words in the word document."""
    
    if __debug__:
        start_time = time.time()

    # Build a formatting table: list of triples [word/space/punctuation: str, underline: bool, bold: bool, highlight: bool]
    formatting_table = []
    
    # Split all_text according to split_string
    text_elements = split_string(all_text)
    
    # Break underlined_text_list and highlighted_text_list according to split_string 
    underlined_elements = []
    for list in underlined_text_list:
        # Check if string contains '...' We need to remove this or it causes underlining errors
        if '...' in list:
            # Delete '...' from the string
            list = list.replace('...', ' ')
        underlined_elements.extend(split_string(list))   
    
    # Bolded elements is already split into 1 word per array slot
    bolded_elements = bolded_text_list
    
    highlighted_elements = []
    for list in highlighted_text_list:
        highlighted_elements.extend(split_string(list))
    # for element in highlighted_elements:
    #     # Remove all non-alphanumeric characters
    #     element = re.sub(r'[^a-zA-Z0-9]', '', element)
    
        
    # Do a pass through to make sure no words in underlined, bolded, highlighted are not present in the original article
    # If they are, it will trip the algorithm
    # Create word dictionary with each word mapping to its indices (possible repeats) in all text
    word_dict = {}
    for index, word in enumerate(text_elements):
        if word in word_dict:
            word_dict[word].append(index)
        else:
            word_dict[word] = [index]
    # Print word_dict to words_debug.txt
    with open('words_debug.txt', 'w', encoding='utf-8') as f:
        f.write(str(word_dict))
    
    # Remove bad characters from words
    indices_to_remove = []
    if '\xa0' in word_dict:
        indices_to_remove.extend(word_dict['\xa0'])
    for index in indices_to_remove:
        text_elements[index] = ' '
    
    # Print words to words_list.txt
    with open('words_list.txt', 'w', encoding='utf-8') as f:
        f.write(str(text_elements))

    # Process bolded_elements
    cleaned_bolded_elements = []
    for word in bolded_elements:
        # If the word ends in a comma, remove the comma
        if word[-1] == ',':
            word = word[:-1]
        # If the word is not in word_dict, skip to the next word
        if word not in word_dict:
            continue

        # Add the processed word to the new list
        cleaned_bolded_elements.append(word)
    # Update the original list with the processed words
    bolded_elements = cleaned_bolded_elements
    
    # Process underlined_elements
    cleaned_underlined_elements = []
    for underlined_index, word in enumerate(underlined_elements):
        if word in word_dict:
            # Check that the highlighted element is both in the word_dict and matches the left and right words
            # match_booleans = []
            # for word_dict_index in word_dict[word]:
            #     match_booleans.append(nearest_word_neighbors_equal(words, underlined_elements, word_dict_index, underlined_index))
            # # Check if any 'True' values in match_booleans
            # if any(match_booleans):
                cleaned_underlined_elements.append(word)

    underlined_elements = cleaned_underlined_elements
            
    # Process highlighted_elements
    cleaned_highlighted_elements = []
    for highlight_index, word in enumerate(highlighted_elements):
        if word in word_dict:
            # Check that the highlighted element is both in the word_dict and matches the left and right words
            # match_booleans = []
            # for word_dict_index in word_dict[word]:
            #     match_booleans.append(nearest_word_neighbors_equal(words, highlighted_elements, word_dict_index, highlight_index))
            # # Check if any 'True' values in match_booleans
            # if any(match_booleans):
                cleaned_highlighted_elements.append(word)
    highlighted_elements = cleaned_highlighted_elements

    # Write highlighted elements to highlighted_debug.txt
    with open('highlighted_debug.txt', 'w') as f:
        for element in highlighted_elements:
            f.write(element + '\n')
    # Write underlined_elements to underlined_debug.txt
    with open('underlined_debug.txt', 'w') as f:
        for element in underlined_elements:
            f.write(element + '\n')
    # Write bolded elements to bolded_debug.txt
    with open('bolded_debug.txt', 'w') as f:
        for element in bolded_elements:
            f.write(element + '\n')

    # Now, we determine which elements in text_elements should be underlined, bolded, and highlighted
    text_pointer = 0
    underlined_pointer = 0
    bolded_pointer = 0
    highlighted_pointer = 0
    
    # Iterate through text_elements, determine if each element should be underlined, bolded, and/or highlighted
    while text_pointer <= (len(text_elements)-1):
        underlined = False
        bolded = False
        highlighted = False
        if text_elements[text_pointer].lower() == '\xa0':
            text_pointer += 1
            continue
        if (underlined_pointer <= (len(underlined_elements)-1)) and (text_elements[text_pointer].lower() == underlined_elements[underlined_pointer].lower()):
            underlined_pointer += 1
            underlined = True
            if (bolded_pointer <= (len(bolded_elements)-1)) and (text_elements[text_pointer].lower() == bolded_elements[bolded_pointer].lower()):
                bolded_pointer += 1
                bolded = True
            if (highlighted_pointer <= (len(highlighted_elements)-1)) and (text_elements[text_pointer].lower() == highlighted_elements[highlighted_pointer].lower()):
                highlighted_pointer += 1
                highlighted = True
            
        formatting_table.append([text_elements[text_pointer], underlined, bolded, highlighted])
        text_pointer += 1

    
    # Now we have the formatting table, we can create the word document    

    # Splice formatting table by removing paragraphs at beginning/end with no underlining - OPTIONAL
    # formatting_table = splice_card_top_bottom(formatting_table)
    
    # Now turn double paragraph breaks into single paragraph breaks
    for index, element in enumerate(formatting_table):
        if ('\n' in element[0]):
            underline, bold, highlight = element[1], element[2], element[3]
            formatting_table[index] = ['\n', underline, bold, highlight]

    # Check if first element in formatting table is new-line
    if '\n' in formatting_table[0][0]:
        # Remove new-line from formatting table
        formatting_table = formatting_table[1:]    
    
    # Print formatting table to debug.txt
    if __debug__:
        with open('debug.txt', 'w') as f:
            for row in formatting_table:
                f.write(str(row) + '\n')
    doc = Document()
    
    # Add the card Tag
    header_runs = []
    tag_header = doc.add_paragraph()
    header_runs.append(tag_header.add_run(tag))
    header_runs[0].bold = True
    header_runs[0].font.name = 'Calibri'
    header_runs[0].font.size = Pt(13)

    author_text = f"{article_info['author']} '{article_info['date']}"
    author = doc.add_paragraph()
    header_runs.append(author.add_run(author_text))
    header_runs[1].bold = True
    header_runs[1].font.name = 'Calibri'
    header_runs[1].font.size = Pt(13)

    # Get today's date in MM-DD-YYYY format
    today = date.today()
    today_str = today.strftime("%m-%d-%Y")

    citation_text = f" [{article_info['author']}; {article_info['qualifications']}; {article_info['date']}; {article_info['title']}; {article_info['publication']}; {article_info['url']}; Accessed {today_str}; cut by AI] *double quotes converted to single quotes"
    header_runs.append(author.add_run(citation_text))
    header_runs[2].bold = False
    header_runs[2].font.name = 'Calibri'
    header_runs[2].font.size = Pt(11)

    runs = []
    run_index = 0
    
    p = doc.add_paragraph()
    
    element_pointer = 0
    element = formatting_table[element_pointer]

    curr_run_underline = element[1]
    curr_run_bold = element[2]
    curr_run_highlight = element[3]
    
    # Create the run as we go. Stop the run when formatting changes
    runs.append(element[0])
    # element_pointer += 1
    
    newPara = False
    while element_pointer <= (len(formatting_table)-2):
        element_pointer += 1
        element = formatting_table[element_pointer]
        if (element[0] == "\n"):
            newPara = True
        # Continue the run until a formatting change
        if (element[1] == curr_run_underline) and (element[2] == curr_run_bold) and (element[3] == curr_run_highlight) and not(newPara):
            runs[run_index] += element[0]
        else:
            add_run_to_paragraph(p, runs[run_index], curr_run_underline, curr_run_bold, curr_run_highlight)
            
            # Start a new run with new formatting
            curr_run_underline = element[1]
            curr_run_bold = element[2]
            curr_run_highlight = element[3]
            
            run_index += 1
            
            # Formatting has changed. Add the run and re-set formatting to new values.
            if newPara:
                p = doc.add_paragraph()
                runs.append("")
                newPara = False
                continue
            
            runs.append(element[0])  


            
    add_run_to_paragraph(p, runs[run_index], curr_run_underline, curr_run_bold, curr_run_highlight)
    
    doc.save(f'{tag}.docx')
    print(f"Document {tag}.docx created successfully")
    if __debug__:
        print(f"Formatting execution time: {time.time() - start_time} seconds")
                

def cut_card(tag: str, url: str) -> None:
    if __debug__:
        start_time = time.time()

    # Scrape article info (author, date, title, body)
    article_info = scrape_article(url)
    if article_info is None:
        return
    # Remove any double-quotes in the article body and replace with single quotes
    if 'body' in article_info:
        article_info['body'] = article_info['body'].replace('"', "'")
    else:
        return None
    
    if __debug__:
        print(f"Scrape article execution time: {time.time() - start_time} seconds")
        start_time = time.time()
        
    if article_info is None:
        print("Failed to scrape article. Exiting...")
        return          
    
    # Pass the extracted information to the card cutting function
    title = article_info['title']
    author = article_info['author']
    date = article_info['date']
    body_text = article_info['body']
    tag_line = tag

    if body_text is None or body_text == "":
        return None
    underlined_text_list, bolded_text_list, highlighted_text_list = llm_cut_article(title, author, date, body_text, tag_line)
    
    if underlined_text_list is None:
        return
    
    # Prepare a dictionary for JSON serialization (excluding the set)
    card_formatting = {
        "underline": underlined_text_list,
        "bold": bolded_text_list,
        "highlight": highlighted_text_list
    }
    
    print(json.dumps(card_formatting, indent=4))
    # Log the output
    logging.info(f'Final card formatting: {json.dumps(card_formatting, indent=4)}')
    
    format_card(article_info, tag, body_text, underlined_text_list, bolded_text_list, highlighted_text_list)
    
    if __debug__:
        print(f"Card cutting execution time: {time.time() - start_time} seconds")
    
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Process argument.")
    parser.add_argument('argument', type=str, help='The argument you want evidence for!')
    parser.add_argument('num_cards', type=str, help='The number of cards you want!')

    args = parser.parse_args()

    sources_and_tags = get_sources(args.argument, args.num_cards)
    
    if sources_and_tags == {}: # re-try once
        sources_and_tags = get_sources(args.argument)
        if sources_and_tags == {}:
            exit
    
    for source, tag in sources_and_tags.items():
        cut_card(tag, source)

    
