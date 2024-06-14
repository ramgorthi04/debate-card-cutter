from docx import Document
from docx.oxml import OxmlElement
from tools import *

# Load the document
doc = Document('../DA.docx')
GPT_MODEL = "gpt-4o-2024-05-13"

def extend(heading_2_groups):
    for h2_group in heading_2_groups:
        for h4_group in h2_group['heading_4_groups']:
            author = h4_group['author']
            highlighted_text = h4_group['highlighted_text']
            
            # Check if highlighted_text and author are non-empty
            if author and highlighted_text:
                heading = h4_group['heading']
                prompt = f"""Given the following argument and evidence taken from an author's article, explain the main points of the argument in the following format:
[Brief re-phrase of the argument] (cite author name in parens)
- [Bullet point] corroborating evidence from author's article
- [Bullet point] corroborating evidence from author's article

Aim for 2-3 bullet points. Be very concise.

Argument: {heading}
Author: {author}
Evidence taken from author's article: {highlighted_text}
"""
                # Get the response from OpenAI API
                explanation = get_gpt_response(prompt, gpt_model=GPT_MODEL, json_mode=False)
                # Add the explanation to the h4_group
                h4_group['explanation'] = explanation
                print(explanation)
                
    return heading_2_groups

def extract_highlighted_text(para):
    highlighted_text = ""
    for run in para.runs:
        if run.font.highlight_color is not None:
            highlighted_text += run.text + " "
    return highlighted_text

def extract_author_and_content(paragraphs):
    author = ""
    content_up_to_bracket = ""
    found_bracket = False

    for para in paragraphs:
        if not found_bracket:
            if '[' in para.text:
                index = para.text.index('[')
                author += para.text[:index].strip()
                content_up_to_bracket += para.text[index:].strip()
                found_bracket = True
            else:
                author += para.text.strip() + " "
        else:
            content_up_to_bracket += para.text.strip() + " "
    
    return author.strip(), content_up_to_bracket.strip()

# Initialize variables
heading_2_groups = []
current_heading_2 = None
current_heading_4_group = None

for para in doc.paragraphs:
    if para.style.name == 'Heading 2':
        # Start a new heading 2 group
        current_heading_2 = {
            'heading': para.text,
            'heading_4_groups': []
        }
        heading_2_groups.append(current_heading_2)
        current_heading_4_group = None  # Reset current_heading_4_group
    elif para.style.name == 'Heading 4':
        # Start a new heading 4 group
        if current_heading_2 is not None:  # Ensure there's a current heading 2
            current_heading_4_group = {
                'heading': para.text,
                'content': [],
                'highlighted_text': "",
                'author': ""
            }
            current_heading_2['heading_4_groups'].append(current_heading_4_group)
    elif current_heading_2 and current_heading_4_group and para.style.name in ['Heading 3', 'Normal']:
        # Add content to the current heading 4 group
        current_heading_4_group['content'].append(para)
        if para.style.name == 'Normal':
            current_heading_4_group['highlighted_text'] += extract_highlighted_text(para)

# Extract author and reformat content
for h2_group in heading_2_groups:
    for h4_group in h2_group['heading_4_groups']:
        author, content_up_to_bracket = extract_author_and_content(h4_group['content'])
        h4_group['author'] = author
        h4_group['content'] = content_up_to_bracket

# Extend the heading_2_groups with explanations
extended_heading_2_groups = extend(heading_2_groups)

# Function to insert a paragraph after another paragraph
def insert_paragraph_after(paragraph, text, style=None):
    new_paragraph = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph)
    new_para = doc.paragraphs[-1]._element.getparent().insert(
        doc.paragraphs[-1]._element.getparent().index(doc.paragraphs[-1]._element) + 1, new_paragraph
    )
    new_run = doc.paragraphs[-1].add_run(text)
    if style:
        doc.paragraphs[-1].style = style
    return new_run

# Function to insert a paragraph before another paragraph
def insert_paragraph_before(paragraph, text, style=None):
    new_paragraph = OxmlElement("w:p")
    paragraph._p.addprevious(new_paragraph)
    new_para = paragraph._element.getparent().insert(
        paragraph._element.getparent().index(paragraph._element), new_paragraph
    )
    new_run = doc.paragraphs[-1].add_run(text)
    if style:
        doc.paragraphs[-1].style = style
    return new_run

# Add the explanations to the Word document
for para in doc.paragraphs:
    for h2_group in extended_heading_2_groups:
        for h4_group in h2_group['heading_4_groups']:
            if para.style.name == 'Heading 4' and para.text == h4_group['heading']:
                # Insert the heading 2 before the Heading 4 group
                insert_paragraph_before(para, '2NC Extensions', style='Heading 2')
                # Insert the explanation after this Heading 4 paragraph
                if 'explanation' in h4_group:
                    insert_paragraph_after(para, h4_group['explanation'], style='Heading 4')
                    # Insert an empty paragraph to ensure spacing
                    insert_paragraph_after(para, "", style=None)

# Save the updated document
doc.save('../DA_extended.docx')
